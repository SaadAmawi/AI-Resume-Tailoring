from openai import OpenAI
from dotenv import load_dotenv
import os 
from docx import Document
import Prompts

doc = Document('Resume.docx')
doc1 = Document('Resume_moded.docx')
Resume1 = ""
Resume2 = ""
def parse_doc(doc,resume):
    for para in doc.paragraphs:
            if len(para.text.strip()) > 0:  # Only include non-empty paragraphs
                resume += "\n" + para.text
    print(resume)
parse_doc(doc,Resume1)
parse_doc(doc1,Resume2)

load_dotenv()

JD = """Unreal Engine Developer Evolvex Technology | Dubai
Evolvex Technology is looking for a talented Unreal Engine Developer to be part of our team and help shape the future of game development! If you are passionate about working on innovative projects, using cutting-edge technologies, and being part of a dynamic team, this opportunity is for you.
In this role, you will actively contribute to game development using Unreal Engine 5. We expect you to have a strong command of Blueprint and C++, along with experience in game mechanics design, optimization, and performance enhancement. If you have the ability to understand and develop complex systems, create innovative solutions to technical challenges, and collaborate effectively within a team, we would love to have you on board.
At Evolvex Technology, we offer more than just a job—we provide an inspiring work environment with career growth opportunities. Working from our modern office in Dubai, you will be part of a dynamic, innovative, and supportive team, contributing to projects that are shaping the future of the gaming industry. You will have the chance to collaborate with industry experts, enhance your skills, and make a significant impact.
"""
JD.replace("\n"," ")





from openai import OpenAI

client = OpenAI(organization=(os.environ['ORGANIZATION_ID']),project=(os.environ['PROJECT']),)


def Tailor_message(Resume, JD):
    completion = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "assistant", "content": Tailor_prompt},
            {"role": "user", "content": f"Here is my Resume: \n ${Resume} and here is the Job Description: \n ${JD}"}
        ]
    )
    responese = completion.choices[0].message.content
    formatted_response = responese.replace('**',' ')
    print(formatted_response)
    return formatted_response



# Tailored_Resume = Tailor_message(Resume,JD)
# Resume_arr = Tailored_Resume.split("\n")
# # print(Resume_arr)

# for i in range(len(doc.paragraphs)):
#     text1 = doc.paragraphs[i].text
#     if len(text1) == 0: # if there is nothing on this line of text
#         continue
#     elif text1=='\n':   # if this line of text is just a blank line
#         continue
#     elif len(text1) > 1 and i <= len(Resume_arr)-1:
#         doc.paragraphs[i].text = Resume_arr[i]
#     else:
#         doc.paragraphs[i].text = doc.paragraphs[i].text
# doc.save('Resume_moded.docx')




def Evaluate_message(Resume, JD):
    completion = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "assistant", "content": Prompts.Evaluate_prompt},
            {"role": "user", "content": f"Here is my Resume: \n ${Resume} and here is the Job Description: \n ${JD}"}
        ]
    )
    responese = completion.choices[0].message.content
    print(responese.replace('\n',' ').replace('**',' ').replace("-","\n"))


print("\n\n-----------------------------------\n\n")
Eval = Evaluate_message(Resume1, JD)
print("\n\n-----------------------------------\n\n")
Eval = Evaluate_message(Resume2, JD)
